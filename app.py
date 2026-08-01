import os
import pandas as pd
import qrcode
import barcode
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask import Flask, render_template, request, send_file
from xhtml2pdf import pisa

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Helper: Word Cell Margins (Padding)
def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Helper: Word Cell Border
def set_cell_border(cell, color="1A73E8", sz="8", val="single"):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val)
        element.set(qn('w:sz'), sz)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)
        tcBorders.append(element)

# Generate Word Document (.docx) - 9cm x 8cm Grid
def build_docx(items):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    chunk_size = 6
    chunks = [items[i:i + chunk_size] for i in range(0, len(items), chunk_size)]

    for page_idx, chunk in enumerate(chunks):
        if page_idx > 0:
            doc.add_page_break()

        grid_table = doc.add_table(rows=3, cols=2)
        grid_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        grid_table.autofit = False

        for i, row in enumerate(chunk):
            row_idx = i // 2
            col_idx = i % 2
            
            grid_table.rows[row_idx].height = Cm(8.0)
            cell = grid_table.cell(row_idx, col_idx)
            cell.width = Cm(9.0)

            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            set_cell_border(cell, color="1A73E8", sz="8")

            fields = [
                ("Model", row.get('Model', 'N/A')),
                ("Serial No", row.get('Serial No', 'N/A')),
                ("Processor Details", row.get('Processor Details', 'N/A')),
                ("Storage", f"{row.get('Harddisk','N/A')} | {row.get('SSD','N/A')} | {row.get('RAM','N/A')}"),
                ("IP No", row.get('IP No', 'N/A')),
                ("Hostname", row.get('Hostname', 'N/A')),
                ("MAC address", row.get('MAC address', 'N/A'))
            ]

            qr_content = "\n".join([f"{l}: {v}" for l, v in fields])
            qr_file = os.path.abspath(f"temp_qr_docx_{page_idx}_{i}.png")
            bar_base = os.path.abspath(f"temp_bar_docx_{page_idx}_{i}")

            qr_img = qrcode.make(qr_content)
            qr_img.save(qr_file)

            code128 = barcode.get_barcode_class('code128')
            serial_barcode = str(row.get('Serial No', '000000'))
            barcode_obj = code128(serial_barcode, writer=ImageWriter())
            barcode_file = barcode_obj.save(bar_base, options={"write_text": False})

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)

            run_title = p.add_run("IT ASSET TAG\n")
            run_title.bold = True
            run_title.font.name = 'Arial'
            run_title.font.size = Pt(10)

            for label, val in fields:
                p_item = cell.add_paragraph()
                p_item.paragraph_format.space_before = Pt(0)
                p_item.paragraph_format.space_after = Pt(1)

                lbl_run = p_item.add_run(f"{label}: ")
                lbl_run.bold = True
                lbl_run.font.name = 'Arial'
                lbl_run.font.size = Pt(8)

                val_run = p_item.add_run(str(val))
                val_run.font.name = 'Arial'
                val_run.font.size = Pt(8)

            bottom_table = cell.add_table(rows=1, cols=2)
            bottom_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            cell_qr = bottom_table.cell(0, 0)
            cell_bar = bottom_table.cell(0, 1)

            p_qr = cell_qr.paragraphs[0]
            p_qr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_qr.add_run().add_picture(qr_file, width=Cm(1.8))

            p_bar = cell_bar.paragraphs[0]
            p_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_bar.add_run().add_picture(barcode_file, width=Cm(3.8), height=Cm(0.9))

            p_bar_txt = cell_bar.add_paragraph()
            p_bar_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            txt_run = p_bar_txt.add_run(f"*{serial_barcode}*")
            txt_run.font.name = 'Arial'
            txt_run.font.size = Pt(7)

            if os.path.exists(qr_file): os.remove(qr_file)
            if os.path.exists(barcode_file): os.remove(barcode_file)

    output_path = os.path.join(OUTPUT_FOLDER, "Asset_Tags.docx")
    doc.save(output_path)
    return output_path

# Generate PDF Document (.pdf) - 9cm x 8cm Layout in Arial
def build_pdf(items):
    temp_images = []
    
    chunk_size = 6
    chunks = [items[i:i + chunk_size] for i in range(0, len(items), chunk_size)]
    
    pages_html = ""

    for page_idx, chunk in enumerate(chunks):
        rows_html = ""
        
        for r in range(0, len(chunk), 2):
            pair = chunk[r:r+2]
            cols_html = ""
            
            for c_idx, row in enumerate(pair):
                idx = page_idx * 6 + r + c_idx
                
                fields = [
                    ("Model", row.get('Model', 'N/A')),
                    ("Serial No", row.get('Serial No', 'N/A')),
                    ("Processor Details", row.get('Processor Details', 'N/A')),
                    ("Storage", f"{row.get('Harddisk','N/A')} | {row.get('SSD','N/A')} | {row.get('RAM','N/A')}"),
                    ("IP No", row.get('IP No', 'N/A')),
                    ("Hostname", row.get('Hostname', 'N/A')),
                    ("MAC address", row.get('MAC address', 'N/A'))
                ]

                qr_content = "\n".join([f"{l}: {v}" for l, v in fields])
                qr_file = os.path.abspath(f"temp_qr_pdf_{idx}.png")
                bar_base = os.path.abspath(f"temp_bar_pdf_{idx}")

                qr_img = qrcode.make(qr_content)
                qr_img.save(qr_file)

                code128 = barcode.get_barcode_class('code128')
                serial_barcode = str(row.get('Serial No', '000000'))
                barcode_obj = code128(serial_barcode, writer=ImageWriter())
                barcode_file = barcode_obj.save(bar_base, options={"write_text": False})

                temp_images.extend([qr_file, barcode_file])

                fields_html = ""
                for label, val in fields:
                    fields_html += f"""
                    <tr>
                        <td class="lbl">{label}:</td>
                        <td class="val">{val}</td>
                    </tr>
                    """

                sticker_inner = f"""
                <div class="sticker-card">
                    <div class="sticker-header">IT ASSET TAG</div>
                    <div class="sticker-content">
                        <table class="info-table">
                            {fields_html}
                        </table>
                        <table class="code-table">
                            <tr>
                                <td style="width: 35%; text-align: left; vertical-align: bottom;">
                                    <img src="{qr_file}" width="52" height="52" />
                                </td>
                                <td style="width: 65%; text-align: center; vertical-align: bottom;">
                                    <img src="{barcode_file}" width="120" height="25" /><br/>
                                    <span class="serial-txt">*{serial_barcode}*</span>
                                </td>
                            </tr>
                        </table>
                    </div>
                </div>
                """
                cols_html += f'<td class="sticker-td">{sticker_inner}</td>'
            
            if len(pair) == 1:
                cols_html += '<td class="sticker-td"></td>'

            rows_html += f'<tr>{cols_html}</tr>'

        page_break_css = 'page-break-after: always;' if page_idx < len(chunks) - 1 else ''
        pages_html += f"""
        <div style="{page_break_css}">
            <table class="page-grid">
                {rows_html}
            </table>
        </div>
        """

    html_full = f"""
    <html>
    <head>
        <style>
            @page {{
                size: A4 portrait;
                margin: 1.2cm 0.8cm 1.2cm 0.8cm;
            }}
            body {{
                font-family: Arial, Helvetica, sans-serif;
                margin: 0;
                padding: 0;
                background-color: #ffffff;
            }}
            .page-grid {{
                width: 100%;
                border-collapse: separate;
                border-spacing: 0.6cm 0.5cm;
            }}
            .sticker-td {{
                width: 9.0cm;
                vertical-align: top;
                padding: 0;
            }}
            .sticker-card {{
                border: 1.5pt solid #1a73e8;
                padding: 0;
                background-color: #ffffff;
                font-family: Arial, Helvetica, sans-serif;
            }}
            .sticker-header {{
                background-color: #1a73e8;
                color: #ffffff;
                font-family: Arial, Helvetica, sans-serif;
                font-weight: bold;
                font-size: 10pt;
                text-align: center;
                padding: 4px 0;
                letter-spacing: 0.5px;
            }}
            .sticker-content {{
                padding: 6px 8px;
            }}
            .info-table {{
                width: 100%;
                border-collapse: collapse;
            }}
            .info-table td {{
                font-family: Arial, Helvetica, sans-serif;
                font-size: 8pt;
                line-height: 1.2;
                padding: 1.5px 0;
                vertical-align: top;
            }}
            .lbl {{
                font-weight: bold;
                color: #000000;
                width: 38%;
            }}
            .val {{
                color: #111111;
                width: 62%;
            }}
            .code-table {{
                width: 100%;
                border-collapse: collapse;
                margin-top: 4px;
            }}
            .serial-txt {{
                font-family: Arial, Helvetica, sans-serif;
                font-size: 7.5pt;
                font-weight: bold;
                color: #222222;
            }}
        </style>
    </head>
    <body>
        {pages_html}
    </body>
    </html>
    """

    output_path = os.path.join(OUTPUT_FOLDER, "Asset_Tags.pdf")
    with open(output_path, "w+b") as out_f:
        pisa.CreatePDF(html_full, dest=out_f)

    for img in temp_images:
        if os.path.exists(img): os.remove(img)

    return output_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    input_type = request.form.get('input_type')
    file_type = request.form.get('file_type')

    items = []

    if input_type == 'excel':
        file = request.files.get('file')
        if not file or file.filename == '':
            return "No excel file selected!"
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(file_path)
        
        df = pd.read_excel(file_path)
        items = df.to_dict('records')

    elif input_type == 'manual':
        models = request.form.getlist('model[]')
        serials = request.form.getlist('serial[]')
        processors = request.form.getlist('processor[]')
        rams = request.form.getlist('ram[]')
        harddisks = request.form.getlist('harddisk[]')
        ssds = request.form.getlist('ssd[]')
        ips = request.form.getlist('ip[]')
        hostnames = request.form.getlist('hostname[]')
        macs = request.form.getlist('mac[]')

        for i in range(len(models)):
            if models[i].strip() or serials[i].strip():
                items.append({
                    'Model': models[i],
                    'Serial No': serials[i],
                    'Processor Details': processors[i],
                    'RAM': rams[i],
                    'Harddisk': harddisks[i],
                    'SSD': ssds[i],
                    'IP No': ips[i],
                    'Hostname': hostnames[i],
                    'MAC address': macs[i]
                })

    if not items:
        return "No data provided!"

    if file_type == 'word':
        out_file = build_docx(items)
        return send_file(out_file, as_attachment=True)
    else:
        out_file = build_pdf(items)
        return send_file(out_file, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
